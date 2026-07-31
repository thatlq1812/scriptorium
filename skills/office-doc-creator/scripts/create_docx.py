#!/usr/bin/env python3
"""Create a .docx from a simple JSON content spec.

Usage:
    python create_docx.py <content.json> <output.docx>

content.json shape:
{
  "title": "Document Title",
  "font": "Times New Roman",
  "blocks": [
    {"type": "heading", "level": 1, "text": "Section 1"},
    {"type": "paragraph", "text": "Some prose.", "font": "Times New Roman"},
    {"type": "table", "headers": ["A", "B"], "rows": [["1", "2"], ["3", "4"]]},
    {"type": "list", "ordering": "ordered", "format": "decimal", "items": ["First", "Second"]}
  ]
}

Font handling
-------------
python-docx's high-level ``style.font.name`` / ``run.font.name`` only sets
the ASCII font slot in the underlying ``w:rFonts`` XML element. Word applies
a *separate* font fallback for East-Asian/complex-script code points
(``w:eastAsia`` / ``w:cs``), and if those are left unset, Vietnamese
diacritics (á, ề, ộ, ư, ơ, đ, ...) can silently render in Word's fallback
theme font instead of the font actually requested. This module sets all
four ``w:rFonts`` attributes (``w:ascii``, ``w:hAnsi``, ``w:cs``,
``w:eastAsia``) explicitly via direct OOXML manipulation, for both
style-level and run-level font requests. See ``metadata.elicited_from`` in
SKILL.md for the source of this fix.

List handling
-------------
Real Word-native numbered/bulleted lists (not string-prefixed fake lists)
via ``docx_numbering.py`` — see that module's docstring.

Both the ``font`` value and the ``list.format`` value are validated against
an explicit allowlist before use; unrecognized values are refused with a
non-zero exit and a clear error rather than silently producing a
plausible-looking but wrong document.
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from docx_numbering import SUPPORTED_LIST_FORMATS, apply_list_item, register_numbering

# Fonts verified present on both Windows and common Word-compatible
# renderers (LibreOffice, Google Docs' Word-import). A font name outside
# this list is refused rather than silently embedded — an unverified font
# name in the XML either gets substituted by whatever Word/LibreOffice
# picks on its own (defeating the point of specifying a font) or, if
# genuinely absent everywhere, renders with missing glyphs. Extend this
# list only after confirming the font is actually available in the target
# rendering environment.
ALLOWED_FONTS: frozenset[str] = frozenset(
    {
        "Times New Roman",
        "Arial",
        "Calibri",
        "Cambria",
        "Courier New",
        "Georgia",
        "Segoe UI",
        "Tahoma",
        "Verdana",
    }
)


def _set_rfonts_xml(rpr, font_name: str) -> None:
    """Set all 4 ``w:rFonts`` attributes on a run-properties (``rPr``) element.

    Removes any pre-existing ``w:rFonts`` child first so repeated calls stay
    idempotent and this function is the sole authority over the font slots.
    """
    for existing in rpr.findall(qn("w:rFonts")):
        rpr.remove(existing)
    rfonts = rpr.makeelement(qn("w:rFonts"), {})
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rpr.insert(0, rfonts)


def _set_style_font(style, font_name: str) -> None:
    """Apply ``font_name`` to a paragraph/character style's run formatting.

    Goes via the underlying XML because ``style.font.name`` alone only sets
    the ASCII slot (see module docstring).
    """
    rpr = style.element.get_or_add_rPr()
    _set_rfonts_xml(rpr, font_name)


def _set_run_font(run, font_name: str) -> None:
    """Apply ``font_name`` to a single run, overriding its style's font."""
    rpr = run._element.get_or_add_rPr()
    _set_rfonts_xml(rpr, font_name)


def _validate_font(font_name: str) -> str:
    if font_name not in ALLOWED_FONTS:
        raise ValueError(
            f"Unknown/unverified font {font_name!r}. Allowed fonts: "
            f"{sorted(ALLOWED_FONTS)}. Add it to ALLOWED_FONTS in "
            f"create_docx.py only after confirming it is actually available "
            f"in the target rendering environment."
        )
    return font_name


def _apply_global_font(doc: Document, font_name: str, heading_levels: set[int]) -> None:
    """Apply ``font_name`` to Normal and every heading style actually used.

    Word's default template sets Heading styles' fonts via theme-font
    references (``w:asciiTheme``) rather than a literal font name, so
    setting Normal's font alone is not guaranteed to cascade to headings —
    each used heading style is set explicitly too.
    """
    _set_style_font(doc.styles["Normal"], font_name)
    for level in heading_levels:
        style_name = "Title" if level == 0 else f"Heading {level}"
        if style_name in doc.styles:
            _set_style_font(doc.styles[style_name], font_name)


def _add_list_block(doc: Document, block: dict) -> None:
    ordering = block.get("ordering")
    if ordering not in ("ordered", "unordered"):
        raise ValueError(
            f"List block 'ordering' must be 'ordered' or 'unordered', got {ordering!r}"
        )

    list_format = block.get("format")
    if list_format not in SUPPORTED_LIST_FORMATS:
        raise ValueError(
            f"Unsupported list format {list_format!r}. Supported: "
            f"{sorted(SUPPORTED_LIST_FORMATS)}"
        )

    items = block.get("items", [])
    if not items:
        raise ValueError("List block must have at least 1 item (got 0)")

    num_id = register_numbering(doc, kind=ordering, list_format=list_format)
    for item_text in items:
        paragraph = doc.add_paragraph(str(item_text))
        apply_list_item(paragraph, num_id, level=0)


def build(content: dict, output_path: Path) -> None:
    doc = Document()

    global_font = content.get("font")
    if global_font is not None:
        _validate_font(global_font)

    blocks = content.get("blocks", [])
    heading_levels = {0} if content.get("title") else set()
    for block in blocks:
        if block.get("type") == "heading":
            heading_levels.add(block.get("level", 1))
    if global_font is not None:
        _apply_global_font(doc, global_font, heading_levels)

    if content.get("title"):
        doc.add_heading(content["title"], level=0)

    for block in blocks:
        btype = block.get("type")
        if btype == "heading":
            doc.add_heading(block["text"], level=block.get("level", 1))
        elif btype == "paragraph":
            block_font = block.get("font")
            if block_font is not None:
                _validate_font(block_font)
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(block["text"])
                _set_run_font(run, block_font)
            else:
                doc.add_paragraph(block["text"])
        elif btype == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            for cell, header in zip(table.rows[0].cells, headers):
                cell.text = str(header)
            for row_data in rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, row_data):
                    cell.text = str(value)
        elif btype == "list":
            _add_list_block(doc, block)
        else:
            raise ValueError(f"Unknown block type: {btype}")

    doc.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("content_json", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()

    content = json.loads(args.content_json.read_text(encoding="utf-8"))
    build(content, args.output_docx)
    print(f"OK: wrote {args.output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
