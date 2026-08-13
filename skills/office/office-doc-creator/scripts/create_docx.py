#!/usr/bin/env python3
"""Create a .docx from a simple JSON content spec.

Usage:
    python create_docx.py <content.json> <output.docx>

content.json shape:
{
  "title": "Document Title",
  "font": "Times New Roman",
  "blocks": [
    {"type": "heading", "level": 1, "text": "Section 1"},
    {"type": "paragraph", "text": "Some prose.", "font": "Times New Roman"},
    {"type": "table", "headers": ["A", "B"], "rows": [["1", "2"], ["3", "4"]]},
    {"type": "list", "ordering": "ordered", "format": "decimal", "items": ["First", "Second"]}
  ]
}

Font handling
-------------
python-docx's high-level ``style.font.name`` / ``run.font.name`` only sets
the ASCII font slot in the underlying ``w:rFonts`` XML element. Word applies
a *separate* font fallback for East-Asian/complex-script code points
(``w:eastAsia`` / ``w:cs``), and if those are left unset, Vietnamese
diacritics (á, ề, ộ, ư, ơ, đ, ...) can silently render in Word's fallback
theme font instead of the font actually requested. This module sets all
four ``w:rFonts`` attributes (``w:ascii``, ``w:hAnsi``, ``w:cs``,
``w:eastAsia``) explicitly via direct OOXML manipulation, for both
style-level and run-level font requests. See ``metadata.elicited_from`` in
SKILL.md for the source of this fix.

List handling
-------------
Real Word-native numbered/bulleted lists (not string-prefixed fake lists)
via ``docx_numbering.py`` — see that module's docstring.

Both the ``font`` value and the ``list.format`` value are validated against
an explicit allowlist before use; unrecognized values are refused with a
non-zero exit and a clear error rather than silently producing a
plausible-looking but wrong document.
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path

import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from docx_numbering import SUPPORTED_LIST_FORMATS, apply_list_item, register_numbering

# Fonts verified present on both Windows and common Word-compatible
# renderers (LibreOffice, Google Docs' Word-import). A font name outside
# this list is refused rather than silently embedded — an unverified font
# name in the XML either gets substituted by whatever Word/LibreOffice
# picks on its own (defeating the point of specifying a font) or, if
# genuinely absent everywhere, renders with missing glyphs. Extend this
# list only after confirming the font is actually available in the target
# rendering environment.
ALLOWED_FONTS: frozenset[str] = frozenset(
    {
        "Times New Roman",
        "Arial",
        "Calibri",
        "Cambria",
        "Courier New",
        "Georgia",
        "Segoe UI",
        "Tahoma",
        "Verdana",
    }
)


def _set_rfonts_xml(rpr, font_name: str) -> None:
    """Set all 4 ``w:rFonts`` attributes on a run-properties (``rPr``) element.

    Removes any pre-existing ``w:rFonts`` child first so repeated calls stay
    idempotent and this function is the sole authority over the font slots.
    """
    for existing in rpr.findall(qn("w:rFonts")):
        rpr.remove(existing)
    rfonts = rpr.makeelement(qn("w:rFonts"), {})
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rpr.insert(0, rfonts)


def _set_style_font(style, font_name: str) -> None:
    """Apply ``font_name`` to a paragraph/character style's run formatting.

    Goes via the underlying XML because ``style.font.name`` alone only sets
    the ASCII slot (see module docstring).
    """
    rpr = style.element.get_or_add_rPr()
    _set_rfonts_xml(rpr, font_name)


def _set_run_font(run, font_name: str) -> None:
    """Apply ``font_name`` to a single run, overriding its style's font."""
    rpr = run._element.get_or_add_rPr()
    _set_rfonts_xml(rpr, font_name)


# Word's own built-in Heading 1-3 theme colors (the "Office 2007-2010" legacy theme
# python-docx's default template ships with) run #365F91 -> #4F81BD -> #4F81BD -- i.e. Heading 2
# and 3 are BOTH a visibly *lighter*, more washed-out blue than Heading 1, not a deeper one
# (confirmed by reading the generated style XML directly, not assumed). A heading hierarchy
# reads backwards when nested sections look paler than their parent. These replacements stay
# dark at every level (owner's own correction, 2026-08-08: "màu text nên là màu xanh đậm đen" --
# text color should be dark navy/near-black, never trending lighter) while still shading
# slightly per level so the hierarchy is visually readable.
DEFAULT_HEADING_COLORS: dict[int, str] = {
    0: "0F2942",  # Title -- near-black navy, darkest
    1: "1B365D",  # Heading 1
    2: "2C4D75",  # Heading 2
    3: "1F3A52",  # Heading 3 and deeper (reused for any level > 3)
}


def _set_style_color(style, color: RGBColor) -> None:
    """Apply an explicit text ``color`` to a paragraph/character style's run formatting.

    Without this, headings render in Word's default theme color (see DEFAULT_HEADING_COLORS'
    docstring above for why that default is a real bug, not a style preference) -- this
    function is what makes the color in this module's output deterministic instead of
    whatever the underlying template's theme happens to define.
    """
    rpr = style.element.get_or_add_rPr()
    for existing in rpr.findall(qn("w:color")):
        rpr.remove(existing)
    color_el = rpr.makeelement(qn("w:color"), {})
    color_el.set(qn("w:val"), str(color))
    rpr.append(color_el)


def _apply_heading_colors(doc: Document, heading_levels: set[int], overrides: dict[int, RGBColor]) -> None:
    """Set an explicit, always-dark color on every Title/Heading-N style actually used.

    `overrides` (from the content spec's top-level `heading_colors`) wins per-level; anything
    not overridden falls back to DEFAULT_HEADING_COLORS, reusing the level-3 shade for any
    level deeper than 3 rather than leaving it unset (which would fall through to the
    lighter-than-its-parent theme default this whole function exists to avoid).
    """
    for level in heading_levels:
        style_name = "Title" if level == 0 else f"Heading {level}"
        if style_name not in doc.styles:
            continue
        color = overrides.get(level) or RGBColor.from_string(
            DEFAULT_HEADING_COLORS.get(level, DEFAULT_HEADING_COLORS[3])
        )
        _set_style_color(doc.styles[style_name], color)


def _validate_font(font_name: str) -> str:
    if font_name not in ALLOWED_FONTS:
        raise ValueError(
            f"Unknown/unverified font {font_name!r}. Allowed fonts: "
            f"{sorted(ALLOWED_FONTS)}. Add it to ALLOWED_FONTS in "
            f"create_docx.py only after confirming it is actually available "
            f"in the target rendering environment."
        )
    return font_name


def _apply_global_font(doc: Document, font_name: str, heading_levels: set[int]) -> None:
    """Apply ``font_name`` to Normal and every heading style actually used.

    Word's default template sets Heading styles' fonts via theme-font
    references (``w:asciiTheme``) rather than a literal font name, so
    setting Normal's font alone is not guaranteed to cascade to headings —
    each used heading style is set explicitly too.
    """
    _set_style_font(doc.styles["Normal"], font_name)
    for level in heading_levels:
        style_name = "Title" if level == 0 else f"Heading {level}"
        if style_name in doc.styles:
            _set_style_font(doc.styles[style_name], font_name)


def _add_list_block(doc: Document, block: dict) -> None:
    ordering = block.get("ordering")
    if ordering not in ("ordered", "unordered"):
        raise ValueError(
            f"List block 'ordering' must be 'ordered' or 'unordered', got {ordering!r}"
        )

    list_format = block.get("format")
    if list_format not in SUPPORTED_LIST_FORMATS:
        raise ValueError(
            f"Unsupported list format {list_format!r}. Supported: "
            f"{sorted(SUPPORTED_LIST_FORMATS)}"
        )

    items = block.get("items", [])
    if not items:
        raise ValueError("List block must have at least 1 item (got 0)")

    num_id = register_numbering(doc, kind=ordering, list_format=list_format)
    for item_text in items:
        paragraph = doc.add_paragraph(str(item_text))
        apply_list_item(paragraph, num_id, level=0)


_HEX_COLOR_RE = re.compile(r"^#?([0-9A-Fa-f]{6})$")


def _validate_hex_color(color: str, field_name: str) -> RGBColor:
    match = _HEX_COLOR_RE.match(color)
    if not match:
        raise ValueError(
            f"{field_name}={color!r} is not a 6-digit hex color (e.g. '1F4E79' or '#1F4E79')."
        )
    return RGBColor.from_string(match.group(1).upper())


def _cover_run(paragraph, text: str, *, size: int, bold: bool = False, italic: bool = False,
                color: RGBColor | None = None, font_name: str | None = None):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if font_name is not None:
        _set_run_font(run, font_name)
    return run


def _add_cover_page_block(doc: Document, block: dict, global_font: str | None) -> None:
    """A formal Vietnamese academic/business cover page (trang bìa) -- institution header,
    document-type label, colored title, a label/value metadata table, closing location/date
    line. Added because the flat heading/paragraph/table/list block set had no way to produce
    one at all: every document this skill built started directly with body content, which reads
    as noticeably plainer than a real thesis/report cover page (see office-doc-creator's own
    "Known limitations": "no company-specific template"). This is not a branded/templated
    solution -- still a JSON content-spec block, same discipline as the rest of the skill, just
    covering the one structural piece (a title page) the spec had no way to express before.
    """
    font_name = block.get("font")
    if font_name is not None:
        _validate_font(font_name)
    else:
        font_name = global_font

    doc_type_color = _validate_hex_color(block.get("document_type_color", "C00000"), "document_type_color")
    title_color = _validate_hex_color(block.get("title_color", "1F4E79"), "title_color")

    # Vertical gaps are real paragraph space_before/space_after (Pt), never a run of empty
    # `add_paragraph()` calls -- besides being the non-idiomatic "press Enter repeatedly" way to
    # space a Word document, stacking many empty paragraphs (each carrying Word's default Normal
    # style line-height + space_after) here produced a genuinely blank extra page between the
    # cover and the real content when rendered (confirmed by rendering to PDF/PNG and looking at
    # it, root-caused by isolating page count across ablation tests, not assumed). One spacer
    # paragraph with an explicit space_after does the same visual job with one paragraph mark
    # instead of N, and is deterministic regardless of the Normal style's own defaults.
    _PT_PER_LINE = 14

    org_lines = block.get("org_lines") or []
    for i, line in enumerate(org_lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        _cover_run(p, str(line), size=13, bold=True, font_name=font_name)

    if block.get("divider", True) and org_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3 * _PT_PER_LINE)
        _cover_run(p, "------------------***------------------", size=13, bold=True, font_name=font_name)

    document_type = block.get("document_type")
    if document_type:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        _cover_run(p, str(document_type), size=15, bold=True, color=doc_type_color, font_name=font_name)

    title = block.get("title")
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        _cover_run(p, str(title), size=16, bold=True, color=title_color, font_name=font_name)

    metadata = block.get("metadata") or []
    if metadata:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(int(block.get("spacing_before_metadata_lines", 6)) * _PT_PER_LINE)
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        # No borders set -- python-docx tables are borderless (no w:tblBorders) by default, which
        # is exactly the "invisible layout grid" look a cover-page label/value block needs.
        # Explicit column widths on both the table and every cell -- `autofit = False` with no
        # width anywhere leaves python-docx writing a `w:tblGrid` with no real column widths, and
        # `table.columns[i].width` alone does not reliably propagate to already-added cells in a
        # multi-row table.
        label_col_width = Cm(5)
        value_col_width = Cm(10)
        table.columns[0].width = label_col_width
        table.columns[1].width = value_col_width
        for entry in metadata:
            row = table.add_row()
            label_cell, value_cell = row.cells
            label_cell.width = label_col_width
            value_cell.width = value_col_width
            label_p = label_cell.paragraphs[0]
            _cover_run(label_p, str(entry.get("label", "")), size=13, bold=True, font_name=font_name)
            value_p = value_cell.paragraphs[0]
            _cover_run(value_p, str(entry.get("value", "")), size=13, italic=True, font_name=font_name)

    footer_line = block.get("footer_line")
    if footer_line:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(int(block.get("spacing_before_footer_lines", 4)) * _PT_PER_LINE)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _cover_run(p, str(footer_line), size=13, bold=True, font_name=font_name)

    if block.get("page_break_after", True):
        doc.add_page_break()


def build(content: dict, output_path: Path) -> None:
    doc = Document()

    global_font = content.get("font")
    if global_font is not None:
        _validate_font(global_font)

    blocks = content.get("blocks", [])
    # A cover_page block renders its own title (possibly styled differently, e.g. colored) as
    # part of the cover layout -- also emitting the top-level "title" as a plain Title-style
    # heading before it would duplicate the document's title on the page. cover_page wins.
    has_cover_page = any(block.get("type") == "cover_page" for block in blocks)
    heading_levels = {0} if content.get("title") and not has_cover_page else set()
    for block in blocks:
        if block.get("type") == "heading":
            heading_levels.add(block.get("level", 1))
    if global_font is not None:
        _apply_global_font(doc, global_font, heading_levels)

    # Heading colors are always set explicitly, not just when a caller opts in -- Word's own
    # default theme is the actual bug being fixed here (DEFAULT_HEADING_COLORS' docstring), so
    # leaving this conditional on `global_font` would silently keep producing the washed-out
    # default for every caller that doesn't ask for a custom font. `heading_colors` (optional,
    # top-level, {"1": "1B365D", ...}) lets a caller pick a different dark palette per level;
    # anything not overridden gets DEFAULT_HEADING_COLORS. Structure (which levels get colored)
    # is enforced; the actual hex values are the caller's call, not a fixed palette -- validated
    # for shape (6-digit hex) only, not restricted to an allowlist like ALLOWED_FONTS is, since
    # an arbitrary color always renders (unlike a font, which can be missing on the renderer).
    raw_heading_colors = content.get("heading_colors") or {}
    heading_color_overrides = {
        int(level): _validate_hex_color(color, f"heading_colors[{level}]")
        for level, color in raw_heading_colors.items()
    }
    _apply_heading_colors(doc, heading_levels, heading_color_overrides)

    if content.get("title") and not has_cover_page:
        doc.add_heading(content["title"], level=0)

    for block in blocks:
        btype = block.get("type")
        if btype == "heading":
            heading_p = doc.add_heading(block["text"], level=block.get("level", 1))
            block_color = block.get("color")
            if block_color is not None:
                color = _validate_hex_color(block_color, "heading block 'color'")
                for run in heading_p.runs:
                    run.font.color.rgb = color
        elif btype == "paragraph":
            block_font = block.get("font")
            if block_font is not None:
                _validate_font(block_font)
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(block["text"])
                _set_run_font(run, block_font)
            else:
                doc.add_paragraph(block["text"])
        elif btype == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            for cell, header in zip(table.rows[0].cells, headers):
                cell.text = str(header)
            for row_data in rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, row_data):
                    cell.text = str(value)
        elif btype == "list":
            _add_list_block(doc, block)
        elif btype == "cover_page":
            _add_cover_page_block(doc, block, global_font)
        else:
            raise ValueError(f"Unknown block type: {btype}")

    doc.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("content_json", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()

    content = json.loads(args.content_json.read_text(encoding="utf-8"))
    build(content, args.output_docx)
    print(f"OK: wrote {args.output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
